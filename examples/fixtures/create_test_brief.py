"""Create the Q2 Model Review test brief — exercises every section type the pipeline handles."""
from pathlib import Path
from docx import Document

doc = Document()

# Title
doc.add_heading("Q2 2025 Credit Model Review — Decision Briefing", level=0)
doc.add_paragraph("For: Credit Risk Committee  |  24 May 2025  |  Prepared by: Credit Risk Modelling Team")

# ── Section 1: Executive Summary ─────────────────────────────────────────────
doc.add_heading("Executive Summary", level=1)
doc.add_paragraph(
    "The Credit Risk Modelling team recommends approval of Credit Model v3 for full production "
    "deployment in Q3 2025. The new model delivers an 18 percentage point improvement in PMSE on "
    "the holdout validation set, passes all regulatory stress tests, and has received independent "
    "model validation sign-off from Accenture Risk Advisory on 14 May 2025. A 90-day parallel run "
    "against live book positions confirms no material change to capital requirements. The primary "
    "residual risk is vintage-level drift in the short-duration cohort, mitigated by a monthly "
    "monitoring protocol agreed with the model risk team."
)
doc.add_paragraph("Management recommendation: Approve Credit Model v3. Target regulatory submission 28 May and go-live August 2025.")

# ── Section 2: Agenda ────────────────────────────────────────────────────────
doc.add_heading("Agenda", level=1)
doc.add_paragraph("Today's committee meeting covers five areas.")
doc.add_paragraph("1. Executive Summary — Recommendation and key findings")
doc.add_paragraph("2. Model Performance — Quantitative results and holdout validation")
doc.add_paragraph("3. Stress Testing and Assumptions — Macro scenario parameters")
doc.add_paragraph("4. Implementation Plan — Go-live timeline and resource requirements")
doc.add_paragraph("5. Appendix — Methodology and supporting data")

# ── Section 3: Q2 Performance Snapshot ───────────────────────────────────────
doc.add_heading("Q2 Performance Snapshot", level=1)
doc.add_paragraph("Six key metrics summarise Q2 2025 model and portfolio performance vs prior quarter.")
doc.add_paragraph("PMSE (holdout, Model v3): 0.046 — down from 0.060 baseline, improvement of 23%")
doc.add_paragraph("Sharpe Ratio: 1.18 — up 0.08 vs Q1 2025")
doc.add_paragraph("AUM covered: $4.7 billion")
doc.add_paragraph("Model coverage: 98.3% of live positions")
doc.add_paragraph("Information Ratio: 0.71")
doc.add_paragraph("Stress test pass rate: 100% across all macro scenarios")

# ── Section 4: Programme Status ───────────────────────────────────────────────
doc.add_heading("Programme Status", level=1)
doc.add_paragraph("Six workstreams are tracked on the path to production deployment.")
doc.add_paragraph("Model Development: Complete. Credit Model v3 signed off by internal validation on 28 April 2025. All factor specifications finalised.")
doc.add_paragraph("Independent Validation: Complete. External validator Accenture confirmed sign-off 14 May 2025 following four-week review.")
doc.add_paragraph("Regulatory Submission: In progress. MRM submission package drafted; awaiting legal review sign-off expected 28 May.")
doc.add_paragraph("Technology Integration: At risk. Two-week delay caused by API versioning incompatibility with core banking system. Vendor patch in testing.")
doc.add_paragraph("Parallel Run: Complete. 90-day parallel run ended 10 May. All cohort results within agreed tolerance bands.")
doc.add_paragraph("Documentation: In progress. Model card and user guide 80% complete. Target sign-off 7 June ahead of go-live.")

# ── Section 5: Model Performance Comparison (table) ──────────────────────────
doc.add_heading("Model Performance Comparison", level=1)
doc.add_paragraph("PMSE by model version and maturity cohort on the out-of-time holdout set (January–March 2025).")

tbl = doc.add_table(rows=5, cols=4)
tbl.style = "Table Grid"
for ci, h in enumerate(["Cohort", "Model v1", "Model v2", "Model v3"]):
    tbl.cell(0, ci).text = h
for ri, row in enumerate([
    ["Short (<1Y)",   "0.078", "0.048", "0.031"],
    ["Medium (1-5Y)", "0.085", "0.061", "0.049"],
    ["Long (>5Y)",    "0.091", "0.072", "0.058"],
    ["All Cohorts",   "0.084", "0.060", "0.046"],
]):
    for ci, val in enumerate(row):
        tbl.cell(ri + 1, ci).text = val

# ── Section 6: Implementation Plan ───────────────────────────────────────────
doc.add_heading("Implementation Plan", level=1)
doc.add_paragraph("The implementation follows a six-milestone schedule from model completion to full production go-live.")
doc.add_paragraph("February 2025: Model development complete — all factor specifications finalised and version-controlled.")
doc.add_paragraph("March 2025: Internal validation sign-off — passed all internal Model Risk Management criteria.")
doc.add_paragraph("May 2025: External validation sign-off — Accenture confirmed; current milestone reached.")
doc.add_paragraph("June 2025: Regulatory submission — MRM package submitted to PRA under SS1/23 requirements.")
doc.add_paragraph("July 2025: Technology integration — API migration, UAT and user acceptance sign-off complete.")
doc.add_paragraph("August 2025: Production go-live — Model v3 active across all credit portfolios.")

# ── Section 7: Stress Test Assumptions (table) ────────────────────────────────
doc.add_heading("Stress Test Assumptions", level=1)
doc.add_paragraph("Model validated against a 1-in-20 adverse macro scenario. Assumptions represent a coherent stress, not individual extremes.")

tbl2 = doc.add_table(rows=7, cols=4)
tbl2.style = "Table Grid"
for ci, h in enumerate(["Assumption", "Value", "Source", "Sensitivity"]):
    tbl2.cell(0, ci).text = h
for ri, row in enumerate([
    ["GDP growth (2025)",   "-1.5%",   "IMF WEO",        "High"],
    ["Rate path (10Y UST)", "+75bps",  "Fed dots",        "High"],
    ["IG spread widening",  "+120bps", "Internal",        "Medium"],
    ["Equity drawdown",     "-20%",    "Internal",        "High"],
    ["FX (EURUSD)",         "0.95",    "Bloomberg",       "Low"],
    ["Recovery rate",       "40%",     "S&P LossStats",   "Medium"],
]):
    for ci, val in enumerate(row):
        tbl2.cell(ri + 1, ci).text = val

# ── Section 8: Factor Return Attribution (table for bar chart) ───────────────
doc.add_heading("Factor Return Attribution", level=1)
doc.add_paragraph("Q2 2025 return attribution by risk factor, comparing Model v2 and Model v3 contributions in basis points.")

tbl3 = doc.add_table(rows=6, cols=3)
tbl3.style = "Table Grid"
for ci, h in enumerate(["Factor", "Model v2 (bps)", "Model v3 (bps)"]):
    tbl3.cell(0, ci).text = h
for ri, row in enumerate([
    ["Equity Beta", "28",  "32"],
    ["Duration",    "-12", "-8"],
    ["Credit",      "15",  "18"],
    ["FX Carry",    "6",   "6"],
    ["Residual",    "-4",  "-2"],
]):
    for ci, val in enumerate(row):
        tbl3.cell(ri + 1, ci).text = val

# ── Section 9: Historical PMSE Trend (table for line chart) ──────────────────
doc.add_heading("Historical PMSE Trend", level=1)
doc.add_paragraph("Rolling quarterly PMSE for Model v2 and Model v3 over the 2023–2025 parallel run period. Lower is better.")

tbl4 = doc.add_table(rows=9, cols=3)
tbl4.style = "Table Grid"
for ci, h in enumerate(["Quarter", "Model v2", "Model v3"]):
    tbl4.cell(0, ci).text = h
for ri, row in enumerate([
    ["Q3 23", "0.065", "0.058"],
    ["Q4 23", "0.063", "0.055"],
    ["Q1 24", "0.061", "0.052"],
    ["Q2 24", "0.060", "0.049"],
    ["Q3 24", "0.062", "0.048"],
    ["Q4 24", "0.061", "0.047"],
    ["Q1 25", "0.060", "0.047"],
    ["Q2 25", "0.060", "0.046"],
]):
    for ci, val in enumerate(row):
        tbl4.cell(ri + 1, ci).text = val

# ── Section 10: Core Finding ─────────────────────────────────────────────────
doc.add_heading("Core Finding", level=1)
doc.add_paragraph(
    "Bills and short-duration instruments drive almost all of the PMSE improvement. "
    "The model's treatment of prepayment risk and vintage effects in the short-duration cohort "
    "explains 72% of the total PMSE gain — equivalent to 14 percentage points of the 18pp total "
    "improvement. The remaining 4pp comes from improved credit spread factor specification in the "
    "medium and long maturity cohorts. External validation using an independent data panel "
    "(January 2020 – December 2024) confirms the finding is robust to different time periods."
)
doc.add_paragraph("Short cohort PMSE reduced by 35%: from 0.048 to 0.031")
doc.add_paragraph("Vintage effect accounts for 72% of the total PMSE gain across all cohorts")
doc.add_paragraph("External validation on independent panel confirms finding holds out-of-sample")
doc.add_paragraph("Stress test capital impact: immaterial — within tolerance agreed with PRA")

# ── Section 11: Implementation Risks ─────────────────────────────────────────
doc.add_heading("Implementation Risk Prioritisation", level=1)
doc.add_paragraph("Four risks require committee attention, assessed by potential impact and resolution urgency.")
doc.add_paragraph("API versioning delay: High impact, High urgency. Two-week delay on technology integration. Extended parallel run covers any gap.")
doc.add_paragraph("Vintage drift monitoring: High impact, Low urgency. Post-launch monthly monitoring protocol agreed. Recalibration triggers if 5% threshold breached.")
doc.add_paragraph("User training coverage: Medium impact, Medium urgency. Three workshops scheduled June. All users trained before go-live.")
doc.add_paragraph("Documentation completeness: Low impact, Low urgency. Model card 80% done. No blocker to go-live; can complete post-launch.")

out = Path("q2_model_review_brief.docx")
doc.save(str(out))
print(f"Created: {out.resolve()}")
