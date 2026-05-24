"""Extract structured source material from Word files for deck generation.

This module does not try to make narrative decisions itself. It gives an AI
assistant a reliable representation of headings, talking points, and tables so
the assistant can choose slide patterns and build a checked presentation.

Designed to tolerate wayward input: bold paragraphs used as headings, international
Word heading styles (Titre, Uberschrift, etc.), N/A and dash cells in tables,
and common unit suffixes like bps, pp, x in numeric columns.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


_DISTIL_THRESHOLD = 120  # paragraphs longer than this need rewriting, not verbatim copying

# Heading style name prefixes across common Word localisations
_HEADING_PREFIXES = (
    "heading",       # English
    "titre",         # French
    "berschrift",    # German (Überschrift — match after stripping non-ASCII)
    "encabezado",    # Spanish
    "intestazione",  # Italian
    "rubrik",        # Swedish / Danish
    "kop ",          # Dutch
    "naglek",        # Polish (Nagłówek — match after stripping)
    "overskrift",    # Norwegian
)

# Values that mean "no data" — skip when assessing whether a column is numeric
_NULL_VALUES = frozenset({
    "", "-", "--", "---", "n/a", "na", "n.a.", "n.m.", "nil",
    "none", "tbd", "tbc", "tba", "—", "–", "?", ".",
})


def _is_heading_style(style_name: str) -> bool:
    """True for any heading style regardless of language."""
    normalised = style_name.lower().encode("ascii", errors="ignore").decode()
    return any(normalised.startswith(prefix) for prefix in _HEADING_PREFIXES)


def _heading_level(style_name: str) -> int:
    """Extract the numeric level from a heading style name (e.g. 'Heading 2' → 2)."""
    match = re.search(r"\d+", style_name)
    return int(match.group()) if match else 1


def _is_implied_heading(paragraph) -> bool:
    """True when a non-heading-styled paragraph looks like a section title.

    Criteria: <= 12 words AND (all runs are bold OR text is ALL CAPS).
    Does not fire on sentences that end with a period.
    """
    text = paragraph.text.strip()
    if not text or text.endswith(".") or len(text.split()) > 12:
        return False
    if text == text.upper() and len(text) > 3:
        return True
    runs = [r for r in paragraph.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def read_source_docx(path: str | Path) -> dict:
    """Read a DOCX source brief into structured paragraphs, sections and tables."""
    source_path = Path(path)
    if not source_path.exists():
        raise FileNotFoundError(f"Source DOCX not found: {source_path}")

    document = Document(str(source_path))
    sections = []
    current = _new_section("Source notes", level=0)
    sections.append(current)
    tables = []

    for block in _iter_body_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            style = block.style.name if block.style else ""
            if _is_heading_style(style):
                current = _new_section(text, level=_heading_level(style))
                sections.append(current)
            elif _is_implied_heading(block):
                current = _new_section(text, level=1, inferred_heading=True)
                sections.append(current)
            else:
                current["paragraphs"].append(text)
                current["blocks"].append({"type": "paragraph", "text": text})
                if len(text) > _DISTIL_THRESHOLD:
                    current["distillation_needed"] = True
        else:
            table = _read_table(block, len(tables) + 1, current["heading"])
            if table is None:
                continue
            tables.append(table)
            current["table_ids"].append(table["id"])
            current["blocks"].append({"type": "table", "table_id": table["id"]})

    non_empty_sections = [
        section for section in sections
        if section["paragraphs"] or section["table_ids"] or section["heading"] != "Source notes"
    ]
    supplemental_text = _supplemental_text(document)
    return {
        "source_path": str(source_path),
        "title": _infer_title(non_empty_sections, source_path),
        "sections": non_empty_sections,
        "tables": tables,
        "supplemental_text": supplemental_text,
        "suggested_visual_inputs": _suggest_visual_inputs(tables),
        "sections_needing_distillation": sum(1 for s in non_empty_sections if s.get("distillation_needed")),
    }


def source_to_markdown(source: dict) -> str:
    """Format structured source input as compact context for an AI assistant."""
    lines = [f"# {source['title']}", "", f"Source: {source['source_path']}"]
    tables_by_id = {table["id"]: table for table in source["tables"]}
    for section in source["sections"]:
        lines.extend(["", f"## {section['heading']}"])
        if section.get("distillation_needed"):
            lines.append("<!-- DISTIL: rewrite long paragraphs below into concise slide copy before using -->")
        blocks = section.get("blocks") or [
            {"type": "paragraph", "text": paragraph} for paragraph in section["paragraphs"]
        ]
        for block in blocks:
            if block["type"] == "paragraph":
                paragraph = block["text"]
                tag = "[DISTIL] " if len(paragraph) > _DISTIL_THRESHOLD else ""
                lines.append(f"- {tag}{paragraph}")
            elif block["type"] == "table":
                _append_table_markdown(lines, tables_by_id[block["table_id"]])
    if source["suggested_visual_inputs"]:
        lines.extend(["", "## Candidate slide visuals"])
        lines.extend(f"- {item}" for item in source["suggested_visual_inputs"])
    for label, texts in source.get("supplemental_text", {}).items():
        if texts:
            lines.extend(["", f"## Captured {label.replace('_', ' ').title()}"])
            lines.extend(f"- {text}" for text in texts)
    return "\n".join(lines) + "\n"


def _new_section(heading: str, level: int, inferred_heading: bool = False) -> dict:
    section = {
        "heading": heading,
        "level": level,
        "paragraphs": [],
        "table_ids": [],
        "blocks": [],
        "distillation_needed": False,
    }
    if inferred_heading:
        section["inferred_heading"] = True
    return section


def _iter_body_blocks(document):
    """Yield top-level body paragraphs and tables in document order."""
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _read_table(table: Table, index: int, section_heading: str) -> dict | None:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return None
    return {
        "id": f"table_{index}",
        "section_heading": section_heading,
        "headers": rows[0],
        "rows": rows[1:],
        "row_count": max(len(rows) - 1, 0),
        "column_count": len(rows[0]),
        "numeric_columns": _numeric_columns(rows),
    }


def _append_table_markdown(lines: list[str], table: dict) -> None:
    lines.extend(["", f"### {table['id']}"])
    lines.append(" | ".join(table["headers"]))
    lines.append(" | ".join("---" for _ in table["headers"]))
    lines.extend(" | ".join(row) for row in table["rows"])
    numeric = ", ".join(table["numeric_columns"]) or "none"
    lines.append(f"Numeric columns detected: {numeric}")


def _supplemental_text(document) -> dict[str, list[str]]:
    """Capture visible source text outside ordinary body paragraphs/tables."""
    headers, footers = [], []
    for section in document.sections:
        headers.extend(_unique_paragraph_text(section.header.paragraphs, headers))
        footers.extend(_unique_paragraph_text(section.footer.paragraphs, footers))

    text_boxes = []
    for text_box in document.element.body.xpath(".//w:txbxContent"):
        paragraphs = []
        for paragraph in text_box.xpath(".//w:p"):
            text = "".join(paragraph.xpath(".//w:t/text()")).strip()
            if text:
                paragraphs.append(text)
        combined = "\n".join(paragraphs)
        if combined and combined not in text_boxes:
            text_boxes.append(combined)

    return {"headers": headers, "footers": footers, "text_boxes": text_boxes}


def _unique_paragraph_text(paragraphs, existing: list[str]) -> list[str]:
    items = []
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text and text not in existing and text not in items:
            items.append(text)
    return items


def _infer_title(sections: list[dict], source_path: Path) -> str:
    for section in sections:
        if section["level"] == 1:
            return section["heading"]
    return source_path.stem.replace("_", " ").title()


def _numeric_columns(rows: list[list[str]]) -> list[str]:
    if len(rows) < 2:
        return []
    headers = rows[0]
    numeric = []
    for index, header in enumerate(headers):
        values = [row[index] for row in rows[1:] if index < len(row)]
        non_null = [v for v in values if not _is_null_value(v)]
        if non_null and sum(_looks_numeric(v) for v in non_null) / len(non_null) >= 0.8:
            numeric.append(header)
    return numeric


def _is_null_value(value: str) -> bool:
    return value.strip().lower() in _NULL_VALUES


def _looks_numeric(value: str) -> bool:
    cleaned = (
        value.replace(",", "")
        .replace("%", "")
        .replace("$", "")
        .replace("(", "-")
        .replace(")", "")
        .lower()
        .replace("bps", "")
        .replace("pp", "")
        .replace("pts", "")
        .replace("pt", "")
        .rstrip("x")  # multiples like 2.1x
        .strip()
    )
    try:
        float(cleaned)
        return True
    except ValueError:
        return False


def _suggest_visual_inputs(tables: list[dict]) -> list[str]:
    suggestions = []
    for table in tables:
        numeric_count = len(table["numeric_columns"])
        if numeric_count >= 2 and table["row_count"] >= 2:
            suggestions.append(
                f"{table['id']}: comparison table, heat map or diverging bars "
                f"using {', '.join(table['numeric_columns'])}"
            )
        elif numeric_count == 1 and table["row_count"] >= 2:
            suggestions.append(
                f"{table['id']}: bar chart or ranked result view using {table['numeric_columns'][0]}"
            )
    return suggestions


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract Word source material for deck generation.")
    parser.add_argument("source_docx", help="Input Word file containing notes and/or tables")
    parser.add_argument("--json", dest="json_path", help="Write structured JSON context")
    parser.add_argument("--markdown", dest="markdown_path", help="Write AI-friendly markdown context")
    args = parser.parse_args()

    source = read_source_docx(args.source_docx)
    if args.json_path:
        target = Path(args.json_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(json.dumps(source, indent=2), encoding="utf-8")
    markdown = source_to_markdown(source)
    if args.markdown_path:
        target = Path(args.markdown_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(markdown, encoding="utf-8")
    else:
        print(markdown)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
