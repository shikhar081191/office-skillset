"""Dependency-light structural quality checks for Word deliverables."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document


PLACEHOLDER_RE = re.compile(r"(lorem ipsum|\bplaceholder\b|\bTODO\b|\[\s*insert\b)", re.IGNORECASE)


def _issue(level: str, check: str, message: str) -> dict:
    return {"level": level, "check": check, "message": message}


def audit_docx(path: str | Path, final: bool = False) -> dict:
    docx_path = Path(path)
    errors = []
    warnings = []
    checks = {"content": 0, "placeholder_text": 0, "heading_order": 0, "table_density": 0}
    if not docx_path.exists():
        errors.append(_issue("error", "file", f"File not found: {docx_path}"))
        return _report(docx_path, final, errors, warnings, checks)
    try:
        document = Document(str(docx_path))
    except Exception as exc:
        errors.append(_issue("error", "file", f"Could not open DOCX: {exc}"))
        return _report(docx_path, final, errors, warnings, checks)

    visible_text = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    visible_text.extend(
        cell.text.strip()
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    )
    checks["content"] = len(visible_text)
    if not visible_text:
        issue = _issue(
            "error" if final else "warning",
            "content",
            "Document contains no visible text.",
        )
        (errors if final else warnings).append(issue)

    for text in visible_text:
        checks["placeholder_text"] += 1
        if PLACEHOLDER_RE.search(text):
            issue = _issue(
                "error" if final else "warning",
                "placeholder_text",
                "Document contains unfinished placeholder text.",
            )
            (errors if final else warnings).append(issue)
            break

    previous_level = 0
    for paragraph in document.paragraphs:
        style = paragraph.style.name if paragraph.style else ""
        if not style.startswith("Heading"):
            continue
        level_text = style.replace("Heading", "").strip()
        if not level_text.isdigit():
            continue
        checks["heading_order"] += 1
        level = int(level_text)
        if previous_level and level > previous_level + 1:
            warnings.append(_issue("warning", "heading_order", f"Heading level jumps from {previous_level} to {level}."))
        previous_level = level

    for table in document.tables:
        checks["table_density"] += 1
        rows = len(table.rows)
        cols = len(table.columns)
        if rows > 30 or cols > 8:
            warnings.append(_issue("warning", "table_density", f"Large table detected ({rows} rows x {cols} columns)."))

    return _report(docx_path, final, errors, warnings, checks)


def _report(path, final, errors, warnings, checks):
    return {
        "path": str(path),
        "final": final,
        "passed": not errors,
        "errors": errors,
        "warnings": warnings,
        "checks": checks,
    }


def ensure_docx_quality(path: str | Path, final: bool = False, report_path: str | Path | None = None) -> dict:
    report = audit_docx(path, final=final)
    if report_path:
        target = Path(report_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(json.dumps(report, indent=2), encoding="utf-8")
    if report["errors"]:
        messages = "; ".join(item["message"] for item in report["errors"])
        raise RuntimeError(f"DOCX quality gate failed: {messages}")
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description="Audit a DOCX before delivery.")
    parser.add_argument("docx_path")
    parser.add_argument("--final", action="store_true")
    parser.add_argument("--report", dest="report_path")
    args = parser.parse_args()
    report = audit_docx(args.docx_path, final=args.final)
    if args.report_path:
        Path(args.report_path).write_text(json.dumps(report, indent=2), encoding="utf-8")
    if report["errors"]:
        for issue in report["errors"]:
            print(f"ERROR: {issue['message']}")
        return 1
    print(f"QA passed: {len(report['warnings'])} warning(s), {args.docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
