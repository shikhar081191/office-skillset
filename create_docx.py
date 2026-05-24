"""Create .docx files using python-docx (pure Python, no npm needed).

This replaces the docx-js / npm workflow for Word document creation.
python-docx covers ~90% of use cases. For advanced features (TOC, tracked
changes, headers/footers), use the unpack/pack XML workflow instead.

Usage examples — run as a script to see a demo:
    python create_docx.py

Or import and use the helpers directly:
    from create_docx import DocxBuilder
    b = DocxBuilder()
    b.heading("My Report", level=1)
    b.paragraph("Some text here.")
    b.save("output.docx")
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class DocxBuilder:
    """Fluent builder for Word documents."""

    def __init__(self, template=None):
        self.doc = Document(template) if template else Document()
        self._set_page_size_letter()

    def _set_page_size_letter(self):
        """Set US Letter page size (python-docx defaults to A4)."""
        from docx.oxml.ns import qn
        from docx.shared import Inches
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def heading(self, text: str, level: int = 1) -> "DocxBuilder":
        """Add a heading (level 1-4)."""
        self.doc.add_heading(text, level=level)
        return self

    def paragraph(self, text: str, bold=False, italic=False,
                  font_size=12, color=None, align="left") -> "DocxBuilder":
        """Add a paragraph with optional formatting."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        if color:
            # color: hex string like "FF0000" or tuple (r, g, b)
            if isinstance(color, str):
                r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
            else:
                r, g, b = color
            run.font.color.rgb = RGBColor(r, g, b)

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        return self

    def bullet_list(self, items: list, style="List Bullet") -> "DocxBuilder":
        """Add a bulleted list."""
        for item in items:
            self.doc.add_paragraph(item, style=style)
        return self

    def numbered_list(self, items: list) -> "DocxBuilder":
        """Add a numbered list."""
        for item in items:
            self.doc.add_paragraph(item, style="List Number")
        return self

    def table(self, data: list[list], header_row=True,
              col_widths=None, header_color="D5E8F0") -> "DocxBuilder":
        """
        Add a table from a 2D list.

        data: list of rows, each row is a list of cell values
        header_row: shade first row as header
        col_widths: list of widths in inches (e.g. [2.0, 1.5, 1.5])
        header_color: hex fill colour for header row
        """
        if not data:
            return self

        rows, cols = len(data), len(data[0])
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        # Set column widths
        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Inches(width)

        # Fill data
        for r_idx, row_data in enumerate(data):
            row = table.rows[r_idx]
            for c_idx, value in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = str(value)

                # Header row styling
                if r_idx == 0 and header_row:
                    # Bold text
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                    # Background fill
                    _set_cell_background(cell, header_color)

        return self

    def page_break(self) -> "DocxBuilder":
        self.doc.add_page_break()
        return self

    def horizontal_rule(self) -> "DocxBuilder":
        """Add a horizontal rule (border under paragraph)."""
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2E75B6")
        pBdr.append(bottom)
        pPr.append(pBdr)
        return self

    def save(self, output_path: str, validate=True, final=False, report_path=None) -> Path:
        """Save the document and run structural QA unless explicitly disabled."""
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(path)
        if validate:
            from docx_qa import ensure_docx_quality
            ensure_docx_quality(path, final=final, report_path=report_path)
        return path


def _set_cell_background(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ── demo ──────────────────────────────────────────────────────────────────────

def demo():
    b = DocxBuilder()

    b.heading("Quarterly Risk Report", level=1)
    b.paragraph("Prepared by Portfolio Risk Modelling | Q2 2025", italic=True)
    b.horizontal_rule()

    b.heading("Executive Summary", level=2)
    b.paragraph(
        "This report summarises key risk factor movements and scenario outcomes "
        "for the quarter ending June 2025."
    )

    b.heading("Key Findings", level=2)
    b.bullet_list([
        "Equity vol increased 12% quarter-on-quarter",
        "Credit spreads widened across IG and HY",
        "FX exposure remains within limits",
    ])

    b.heading("Factor Performance", level=2)
    b.table(
        data=[
            ["Factor", "Return (%)", "Vol (%)", "Sharpe"],
            ["Equity Beta", "3.2", "14.1", "0.23"],
            ["Credit Spread", "-1.8", "6.4", "-0.28"],
            ["Rates Duration", "0.9", "4.2", "0.21"],
            ["FX Carry", "2.1", "8.7", "0.24"],
        ],
        header_row=True,
        col_widths=[2.5, 1.5, 1.5, 1.5],
    )

    b.page_break()
    b.heading("Notes", level=2)
    b.paragraph("All figures are in USD unless otherwise stated.")

    out = b.save("demo_output.docx")
    print(f"Created: {out}")


if __name__ == "__main__":
    demo()
